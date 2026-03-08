"""
文件处理服务 - 支持多模态文件内容提取

v3.7 新增功能：
- PDF 文本提取
- TXT 智能编码检测
- 图片 OCR 文字识别（可选）
"""

from pathlib import Path
from typing import Any, Dict

import aiofiles
import chardet
from loguru import logger


class FileProcessor:
    """多模态文件处理服务"""

    def __init__(
        self, upload_dir: str = "./data/uploads", enable_vision_api: bool = True, vision_provider: str = "openai"
    ):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.enable_vision_api = enable_vision_api
        self.vision_provider = vision_provider

        # 初始化Vision LLM（可选）
        self.vision_llm = None
        if enable_vision_api:
            try:
                from intelligent_project_analyzer.settings import settings

                if vision_provider == "gemini":
                    #  使用 Google Gemini Vision
                    import os

                    from langchain_google_genai import ChatGoogleGenerativeAI

                    api_key = os.getenv("GOOGLE_API_KEY") or settings.llm.api_key
                    self.vision_llm = ChatGoogleGenerativeAI(
                        model="gemini-1.5-flash", temperature=0.7, google_api_key=api_key  # 或 gemini-1.5-pro
                    )
                    logger.info(" Vision API已启用 (Google Gemini)")

                elif vision_provider == "gemini-openrouter":
                    #  通过 OpenRouter 使用 Gemini Vision (解决国内网络限制)
                    # 注意: 目前OpenRouter中Gemini Vision支持有限,推荐使用 openai-openrouter
                    import os

                    from langchain_openai import ChatOpenAI

                    api_key = os.getenv("OPENROUTER_API_KEY")
                    self.vision_llm = ChatOpenAI(
                        model="google/gemini-pro-1.5",  # OpenRouter中的Gemini模型
                        temperature=0.7,
                        api_key=api_key,
                        base_url="https://openrouter.ai/api/v1",
                    )
                    logger.info(" Vision API已启用 (Gemini via OpenRouter - 国内可用)")

                elif vision_provider == "openai-openrouter":
                    #  通过 OpenRouter 使用 GPT-4o Vision (推荐国内用户)
                    import os

                    from langchain_openai import ChatOpenAI

                    api_key = os.getenv("OPENROUTER_API_KEY")
                    self.vision_llm = ChatOpenAI(
                        model="openai/gpt-4o",  # OpenRouter中的GPT-4o
                        temperature=0.7,
                        api_key=api_key,
                        base_url="https://openrouter.ai/api/v1",
                    )
                    logger.info(" Vision API已启用 (GPT-4o via OpenRouter - 国内可用)")

                else:
                    # OpenAI GPT-4 Vision
                    from langchain_openai import ChatOpenAI

                    self.vision_llm = ChatOpenAI(
                        model="gpt-4o", temperature=0.7, api_key=settings.llm.api_key, base_url=settings.llm.api_base
                    )
                    logger.info(" Vision API已启用 (OpenAI GPT-4V)")

            except Exception as e:
                logger.warning(f"️ Vision API初始化失败: {e}")
                self.enable_vision_api = False

        logger.info(f" 文件处理器初始化: {self.upload_dir}")

    # ========================================================================
    #  v7.156: 按需加载图片 base64 - 性能优化
    # ========================================================================

    def load_image_base64(self, file_path: str, max_size: int = 1024) -> str | None:
        """
         v7.156: 按需加载图片为 base64 格式

        只在需要时才加载原图（如概念图生成需要参考原图），
        避免在工作流状态中存储大量图片数据。

        Args:
            file_path: 图片文件路径（绝对路径或相对路径）
            max_size: 最大边长（像素），超过则缩放，默认1024

        Returns:
            base64编码的图片字符串，失败返回None
        """
        import base64
        from pathlib import Path

        try:
            # 处理相对路径
            path = Path(file_path)
            if not path.is_absolute():
                path = self.upload_dir / file_path

            if not path.exists():
                logger.warning(f"️ [v7.156] 图片文件不存在: {path}")
                return None

            import io

            from PIL import Image

            img = Image.open(path)

            # 如果图片过大，按比例缩放（节省token和传输）
            if max(img.size) > max_size:
                ratio = max_size / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
                logger.info(f"️ [v7.156] 图片已缩放: {img.size[0]}x{img.size[1]}")

            # 转换为JPEG格式（更小的体积）
            buffer = io.BytesIO()
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(buffer, format="JPEG", quality=85)

            img_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
            logger.info(f" [v7.156] 图片已加载: {path.name} ({len(img_base64)} chars)")

            return img_base64

        except Exception as e:
            logger.error(f" [v7.156] 加载图片失败: {file_path} - {e}")
            return None

    def resolve_image_path(self, visual_ref: Dict[str, Any]) -> str | None:
        """
         v7.156: 解析视觉参考中的图片路径

        优先使用绝对路径（本地访问快），
        如果绝对路径失效则尝试相对路径（容器/部署兼容）。

        Args:
            visual_ref: 视觉参考字典（包含 file_path 和 relative_path）

        Returns:
            有效的图片路径，无效返回None
        """
        from pathlib import Path

        # 优先尝试绝对路径
        abs_path = visual_ref.get("file_path")
        if abs_path and Path(abs_path).exists():
            return abs_path

        # 尝试相对路径
        rel_path = visual_ref.get("relative_path")
        if rel_path:
            full_path = self.upload_dir / rel_path
            if full_path.exists():
                logger.info(f" [v7.156] 使用相对路径: {rel_path}")
                return str(full_path)

        logger.warning(f"️ [v7.156] 无法解析图片路径: abs={abs_path}, rel={rel_path}")
        return None

    async def save_file(self, file_content: bytes, filename: str, session_id: str) -> Path:
        """
        保存上传的文件

        Args:
            file_content: 文件内容（字节）
            filename: 原始文件名
            session_id: 会话ID

        Returns:
            保存后的文件路径
        """
        # 创建会话专属目录
        session_dir = self.upload_dir / session_id
        session_dir.mkdir(parents=True, exist_ok=True)

        # 生成安全的文件路径
        safe_filename = self._sanitize_filename(filename)
        file_path = session_dir / safe_filename

        # 异步写入文件
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_content)

        logger.info(f" 文件已保存: {file_path} ({len(file_content)} bytes)")
        return file_path

    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名，移除不安全字符"""
        import re

        # 保留中文、字母、数字、点、下划线、连字符
        safe_name = re.sub(r"[^\w\u4e00-\u9fa5\.\-]", "_", filename)
        return safe_name

    async def extract_content(self, file_path: Path, content_type: str) -> Dict[str, Any]:
        """
        根据文件类型提取内容

        Args:
            file_path: 文件路径
            content_type: MIME类型

        Returns:
            提取结果字典
        """
        logger.info(f" 开始提取文件内容: {file_path.name} ({content_type})")

        try:
            if content_type == "application/pdf":
                return await self._extract_pdf(file_path)

            elif content_type == "text/plain":
                return await self._extract_text(file_path)

            elif content_type in ["image/png", "image/jpeg", "image/jpg", "image/webp"]:
                return await self._extract_image(file_path)

            elif (
                "word" in content_type
                or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                return await self._extract_word(file_path)

            elif (
                "excel" in content_type
                or "spreadsheet" in content_type
                or content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ):
                return await self._extract_excel(file_path)

            else:
                return {"type": "unknown", "text": "", "error": f"不支持的文件类型: {content_type}"}

        except Exception as e:
            logger.error(f" 文件提取失败: {file_path.name} - {str(e)}")
            return {"type": "error", "text": "", "error": f"文件处理失败: {str(e)}"}

    async def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        提取PDF内容（文本）

        使用 pdfplumber 提取文本内容
        """
        try:
            import pdfplumber
        except ImportError:
            logger.warning("️ pdfplumber 未安装，尝试使用 PyPDF2")
            return await self._extract_pdf_pypdf2(file_path)

        text_content = []
        total_pages = 0

        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(f"=== 第 {page_num} 页 ===\n{text}")

            full_text = "\n\n".join(text_content)

            logger.info(f" PDF提取完成: {total_pages}页, {len(full_text)}字符")

            return {
                "type": "pdf",
                "text": full_text,
                "pages": total_pages,
                "summary": f"PDF文档，共{total_pages}页，提取文本{len(full_text)}字符",
            }

        except Exception as e:
            logger.error(f" PDF提取失败: {str(e)}")
            return {"type": "pdf", "text": "", "error": f"PDF处理失败: {str(e)}"}

    async def _extract_pdf_pypdf2(self, file_path: Path) -> Dict[str, Any]:
        """
        使用 PyPDF2 提取PDF内容（fallback）
        """
        try:
            import PyPDF2

            text_content = []

            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                total_pages = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(f"=== 第 {page_num} 页 ===\n{text}")

            full_text = "\n\n".join(text_content)

            return {
                "type": "pdf",
                "text": full_text,
                "pages": total_pages,
                "summary": f"PDF文档，共{total_pages}页，提取文本{len(full_text)}字符",
            }

        except Exception as e:
            return {"type": "pdf", "text": "", "error": f"PDF处理失败: {str(e)}"}

    async def _extract_text(self, file_path: Path) -> Dict[str, Any]:
        """
        提取TXT内容（智能编码检测）
        """
        # 读取文件字节
        async with aiofiles.open(file_path, "rb") as f:
            raw_data = await f.read()

        # 使用 chardet 检测编码
        detected = chardet.detect(raw_data)
        encoding = detected.get("encoding", "utf-8")
        confidence = detected.get("confidence", 0)

        logger.info(f" 检测到编码: {encoding} (置信度: {confidence:.2%})")

        # 尝试解码
        encodings_to_try = [encoding, "utf-8", "gbk", "gb2312", "utf-16", "latin-1"]

        for enc in encodings_to_try:
            try:
                content = raw_data.decode(enc)
                logger.info(f" 成功使用编码 {enc} 解码文本文件")

                return {
                    "type": "text",
                    "text": content,
                    "encoding": enc,
                    "length": len(content),
                    "summary": f"文本文件，编码{enc}，共{len(content)}字符",
                }
            except (UnicodeDecodeError, AttributeError):
                continue

        logger.error(" 无法识别文本编码")
        return {"type": "text", "text": "", "error": "无法识别文本编码"}

    async def _extract_image(self, file_path: Path) -> Dict[str, Any]:
        """
        提取图片内容

         v3.8新增: Vision API分析图片内容
        """
        try:
            import base64

            from PIL import Image

            img = Image.open(file_path)
            width, height = img.size
            format_name = img.format

            logger.info(f"️ 图片信息: {width}x{height}, {format_name}")

            #  使用Vision API分析图片内容
            vision_analysis = ""
            if self.enable_vision_api and self.vision_llm:
                try:
                    # 读取图片为base64
                    with open(file_path, "rb") as image_file:
                        image_data = base64.b64encode(image_file.read()).decode()

                    # 构造Vision API请求
                    vision_prompt = """请详细分析这张图片的内容，特别关注设计相关的元素：

1. **主要内容**：描述图片中的主要对象、场景或设计
2. **风格特征**：色彩、材质、风格类型（现代/简约/北欧等）
3. **空间布局**：如果是室内设计，描述空间划分和动线
4. **设计亮点**：值得借鉴的设计细节
5. **文字信息**：图片中包含的任何文字或标识

请用中文回答，聚焦于设计和空间规划相关的信息。"""

                    # 调用Vision API
                    import asyncio

                    from langchain_core.messages import HumanMessage

                    message = HumanMessage(
                        content=[
                            {"type": "text", "text": vision_prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}},
                        ]
                    )

                    #  添加30秒超时保护
                    try:
                        response = await asyncio.wait_for(
                            asyncio.to_thread(self.vision_llm.invoke, [message]), timeout=30.0
                        )
                        vision_analysis = response.content
                        logger.info(f" Vision API分析完成: {len(vision_analysis)}字符")
                    except asyncio.TimeoutError:
                        logger.warning("️ Vision API调用超时（30秒）")
                        vision_analysis = "[Vision API调用超时，跳过图片分析]"

                except Exception as e:
                    logger.warning(f"️ Vision API调用失败: {e}")
                    vision_analysis = f"[Vision API分析失败: {str(e)}]"

            # 构造返回文本
            if vision_analysis:
                text_content = f"""[图片文件: {file_path.name}]
尺寸: {width}x{height}
格式: {format_name}

## AI视觉分析

{vision_analysis}
"""
            else:
                text_content = f"[图片文件: {file_path.name}, 尺寸 {width}x{height}, 格式 {format_name}]"

            return {
                "type": "image",
                "text": text_content,
                "width": width,
                "height": height,
                "format": format_name,
                "vision_analysis": vision_analysis,
                "summary": f"图片文件，{width}x{height}，格式{format_name}" + (" (已AI分析)" if vision_analysis else ""),
            }

        except Exception as e:
            return {"type": "image", "text": f"[图片文件: {file_path.name}]", "error": f"图片处理失败: {str(e)}"}

    # ========================================================================
    #  v7.155: 增强版图片提取 - 支持结构化视觉特征
    # ========================================================================

    async def extract_image_enhanced(self, file_path: Path) -> Dict[str, Any]:
        """
         v7.155: 增强版图片提取 - 提取结构化视觉特征

        相比 _extract_image() 的改进：
        1. 提取结构化视觉特征（颜色、风格、材质等）
        2. 不存储 base64 数据（性能优化，按需加载）
        3. 支持用户追加描述

        Args:
            file_path: 图片文件路径

        Returns:
            包含结构化视觉特征的字典
        """
        try:

            from PIL import Image

            img = Image.open(file_path)
            width, height = img.size
            format_name = img.format

            logger.info(f"️ [v7.155] 增强版图片提取: {file_path.name} ({width}x{height})")

            # 调用基础Vision分析
            basic_result = await self._extract_image(file_path)
            vision_analysis = basic_result.get("vision_analysis", "")

            # 提取结构化视觉特征
            structured_features = await self._extract_structured_visual_features(file_path)

            return {
                "type": "image",
                "file_path": str(file_path),
                "width": width,
                "height": height,
                "format": format_name,
                "vision_analysis": vision_analysis,
                "structured_features": structured_features,
                "user_description": None,
                "reference_type": "general",
                "text": basic_result.get("text", ""),
                "summary": basic_result.get("summary", ""),
            }

        except Exception as e:
            logger.error(f" [v7.155] 增强版图片提取失败: {e}")
            return {
                "type": "image",
                "file_path": str(file_path),
                "error": f"图片处理失败: {str(e)}",
                "structured_features": self._get_default_structured_features(),
            }

    async def _extract_structured_visual_features(self, file_path: Path) -> Dict[str, Any]:
        """
         v7.155: 使用Vision API提取结构化视觉特征

        输出JSON格式的结构化特征，便于后续流程使用

        Args:
            file_path: 图片文件路径

        Returns:
            结构化视觉特征字典
        """
        if not self.enable_vision_api or not self.vision_llm:
            logger.warning("️ Vision API未启用，返回默认特征")
            return self._get_default_structured_features()

        try:
            import asyncio
            import base64
            import json

            from langchain_core.messages import HumanMessage

            # 读取图片为base64
            with open(file_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode()

            structured_prompt = """请分析这张图片并以JSON格式输出以下结构化信息（用于设计参考）：

{
    "dominant_colors": ["主色1", "主色2", "主色3"],
    "style_keywords": ["风格关键词1", "风格关键词2", "风格关键词3"],
    "materials": ["材质1", "材质2"],
    "spatial_layout": "空间布局描述（一句话）",
    "mood_atmosphere": "氛围描述（一句话）",
    "design_elements": ["设计元素1", "设计元素2", "设计元素3"]
}

要求：
1. dominant_colors: 提取3个主要颜色（如"暖白色"、"原木色"、"深灰色"）
2. style_keywords: 提取3-5个风格关键词（如"北欧简约"、"现代工业"、"日式禅意"）
3. materials: 提取2-4个主要材质（如"大理石"、"实木"、"金属"）
4. spatial_layout: 描述空间布局特点（如"开放式客餐厅一体化"）
5. mood_atmosphere: 描述整体氛围（如"温馨舒适的家庭氛围"）
6. design_elements: 提取3-5个设计亮点（如"落地窗"、"悬浮电视柜"、"隐藏式灯带"）

请确保输出有效的JSON格式，不要包含其他文字。"""

            message = HumanMessage(
                content=[
                    {"type": "text", "text": structured_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}},
                ]
            )

            # 调用Vision API（30秒超时）
            response = await asyncio.wait_for(asyncio.to_thread(self.vision_llm.invoke, [message]), timeout=30.0)

            # 解析JSON响应
            content = response.content.strip()

            # 处理可能的markdown代码块
            if "```" in content:
                # 提取代码块内容
                import re

                json_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", content)
                if json_match:
                    content = json_match.group(1).strip()

            features = json.loads(content)
            logger.info(f" [v7.155] 结构化视觉特征提取成功: {list(features.keys())}")
            return features

        except asyncio.TimeoutError:
            logger.warning("️ [v7.155] Vision API调用超时（30秒）")
            return self._get_default_structured_features()
        except json.JSONDecodeError as e:
            logger.warning(f"️ [v7.155] JSON解析失败: {e}")
            return self._get_default_structured_features()
        except Exception as e:
            logger.warning(f"️ [v7.155] 结构化特征提取失败: {e}")
            return self._get_default_structured_features()

    def _get_default_structured_features(self) -> Dict[str, Any]:
        """返回默认的结构化特征"""
        return {
            "dominant_colors": [],
            "style_keywords": [],
            "materials": [],
            "spatial_layout": "",
            "mood_atmosphere": "",
            "design_elements": [],
        }

    # ========================================================================
    #  v7.157: 按类型提取特征
    # ========================================================================

    async def extract_category_specific_features(
        self, file_path: Path, categories: list[str], custom_description: str = ""
    ) -> Dict[str, Any]:
        """
         v7.157: 根据用户选择的分类提取针对性特征

        根据用户选择的参考类型（如色彩、风格、材质等），
        构建针对性的Vision API提示词，提取该类型的特定特征。

        Args:
            file_path: 图片文件路径
            categories: 用户选择的分类列表（如 ["color", "style"]）
            custom_description: 用户自定义描述

        Returns:
            按类型组织的特征字典
        """
        if not self.enable_vision_api or not self.vision_llm:
            logger.warning("️ [v7.157] Vision API未启用，返回空特征")
            return {}

        if not categories:
            logger.info("ℹ️ [v7.157] 未选择分类，跳过针对性特征提取")
            return {}

        # 分类到提示词的映射
        category_prompts = {
            "color": "色彩分析：提取主色调、配色方案、色彩比例、色彩情绪",
            "style": "风格分析：识别设计风格（如北欧/日式/工业/现代等）、时代特征、文化元素",
            "material": "材质分析：识别主要材质（如大理石/实木/金属/布艺等）、表面质感、工艺细节",
            "layout": "布局分析：描述空间分区、动线规划、功能布局、空间比例",
            "environment": "环境分析：描述周边环境、景观元素、建筑外观、自然光线",
            "mood": "氛围分析：描述整体格调、情感表达、生活方式暗示",
            "item": "单品分析：识别家具/灯具/软装单品的类型、造型特征、搭配方式",
            "lighting": "光线分析：描述光源类型、光影效果、照明层次、明暗对比",
        }

        try:
            import asyncio
            import base64
            import json

            from langchain_core.messages import HumanMessage

            # 读取图片为base64
            with open(file_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode()

            # 构建针对性提示词
            selected_prompts = [category_prompts.get(cat, "") for cat in categories if cat in category_prompts]
            if not selected_prompts:
                return {}

            prompt_text = f"""请分析这张图片，重点关注以下方面：

{chr(10).join(f"- {p}" for p in selected_prompts)}

{"用户补充说明: " + custom_description if custom_description else ""}

请以JSON格式输出分析结果，每个分类作为一个键：
{{
{chr(10).join(f'    "{cat}": "该分类的分析结果"' for cat in categories if cat in category_prompts)}
}}

请确保输出有效的JSON格式，不要包含其他文字。"""

            message = HumanMessage(
                content=[
                    {"type": "text", "text": prompt_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}},
                ]
            )

            # 调用Vision API（30秒超时）
            response = await asyncio.wait_for(asyncio.to_thread(self.vision_llm.invoke, [message]), timeout=30.0)

            # 解析JSON响应
            content = response.content.strip()

            # 处理可能的markdown代码块
            if "```" in content:
                import re

                json_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", content)
                if json_match:
                    content = json_match.group(1).strip()

            features = json.loads(content)
            logger.info(f" [v7.157] 按类型特征提取成功: {list(features.keys())}")
            return features

        except asyncio.TimeoutError:
            logger.warning("️ [v7.157] Vision API调用超时（30秒）")
            return {}
        except json.JSONDecodeError as e:
            logger.warning(f"️ [v7.157] JSON解析失败: {e}")
            return {}
        except Exception as e:
            logger.warning(f"️ [v7.157] 按类型特征提取失败: {e}")
            return {}

    @staticmethod
    def load_image_base64(file_path: str) -> str:
        """
         v7.155: 按需加载图片的base64数据

        性能优化：base64数据不存储在state中，需要时按需加载

        Args:
            file_path: 图片文件路径

        Returns:
            Base64编码的图片数据
        """
        import base64
        from pathlib import Path

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"图片文件不存在: {file_path}")

        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()

    async def _extract_word(self, file_path: Path) -> Dict[str, Any]:
        """
        提取Word文档内容 (.docx)

         v3.8新增: 支持Word文档
        """
        try:
            from docx import Document

            doc = Document(file_path)

            # 提取段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 提取表格
            tables_text = []
            for table_idx, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))

                if table_data:
                    tables_text.append(f"\n[表格 {table_idx}]\n" + "\n".join(table_data))

            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n".join(tables_text)

            logger.info(f" Word提取完成: {len(paragraphs)}段落, {len(doc.tables)}表格")

            return {
                "type": "word",
                "text": full_text,
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "summary": f"Word文档，{len(paragraphs)}段落，{len(doc.tables)}表格",
            }

        except Exception as e:
            logger.error(f" Word提取失败: {str(e)}")
            return {"type": "word", "text": "", "error": f"Word处理失败: {str(e)}"}

    async def _extract_excel(self, file_path: Path) -> Dict[str, Any]:
        """
        提取Excel表格内容 (.xlsx)

         v3.8新增: 支持Excel表格
        """
        try:
            import pandas as pd

            # 读取所有工作表
            dfs = pd.read_excel(file_path, sheet_name=None)

            text_parts = []
            total_rows = 0

            for sheet_name, df in dfs.items():
                total_rows += len(df)

                text_parts.append(f"\n=== 工作表: {sheet_name} ({len(df)}行 x {len(df.columns)}列) ===\n")

                # 转换为文本表格（限制前100行）
                if len(df) > 100:
                    text_parts.append(df.head(100).to_string(index=False))
                    text_parts.append(f"\n... (共{len(df)}行，仅显示前100行) ...\n")
                else:
                    text_parts.append(df.to_string(index=False))

            full_text = "\n".join(text_parts)

            logger.info(f" Excel提取完成: {len(dfs)}工作表, {total_rows}行")

            return {
                "type": "excel",
                "text": full_text,
                "sheets": len(dfs),
                "total_rows": total_rows,
                "summary": f"Excel表格，{len(dfs)}工作表，共{total_rows}行",
            }

        except Exception as e:
            logger.error(f" Excel提取失败: {str(e)}")
            return {"type": "excel", "text": "", "error": f"Excel处理失败: {str(e)}"}


def build_combined_input(user_text: str, file_contents: list[Dict[str, Any]]) -> str:
    """
    合并用户输入和文件内容，生成统一的分析输入

    Args:
        user_text: 用户输入的文本
        file_contents: 文件提取结果列表

    Returns:
        合并后的完整输入文本
    """
    parts = []

    # 1. 用户原始输入
    if user_text.strip():
        parts.append(f"## 用户需求描述\n\n{user_text}\n")

    # 2. 附件内容
    if file_contents:
        parts.append("## 附件材料\n")
        parts.append("**说明**: 以下附件为用户提供的背景资料和参考信息，仅供参考。请根据「用户需求描述」中的明确要求，结合这些背景资料进行分析，而不是将所有附件内容都视为必须实现的需求。\n")

        for idx, content in enumerate(file_contents, 1):
            file_type = content.get("type", "unknown")
            text = content.get("text", "")
            summary = content.get("summary", "")
            error = content.get("error", "")

            if error:
                parts.append(f"### 附件 {idx} ({file_type.upper()}) - 处理失败\n\n{error}\n")
                continue

            # 智能截断（保留重要信息）
            if len(text) > 5000:
                text = text[:5000] + f"\n\n...(内容过长，已截断，完整内容共{len(text)}字符)..."

            parts.append(f"### 附件 {idx} ({file_type.upper()})\n")
            if summary:
                parts.append(f"**摘要**: {summary}\n")
            parts.append(f"\n{text}\n")

    combined = "\n".join(parts)

    logger.info(f" 内容合并完成: 原始文本={len(user_text)}字符, 文件={len(file_contents)}个, 合并后={len(combined)}字符")

    return combined


# 全局实例
#  可通过环境变量 VISION_PROVIDER 选择：openai 或 gemini
import os

vision_provider = os.getenv("VISION_PROVIDER", "openai")  # 默认 OpenAI
file_processor = FileProcessor(vision_provider=vision_provider)
